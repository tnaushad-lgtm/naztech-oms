# -*- coding: utf-8 -*-
"""Generate the Naztech OMS User Manual (.docx) with embedded screenshots."""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, "img")
OUT = os.path.join(HERE, "Naztech_OMS_User_Manual.docx")

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
VIOLET = RGBColor(0x6D, 0x28, 0xD9)
CYAN   = RGBColor(0x0E, 0x74, 0x90)
INK    = RGBColor(0x1F, 0x29, 0x37)
MUTE   = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x15, 0x80, 0x3D)
RED    = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.06

_HSP = {1: (12, 3), 2: (8, 2), 3: (6, 2)}
for lvl, color, size in [(1, INDIGO, 17), (2, VIOLET, 13.5), (3, CYAN, 11.5)]:
    st = doc.styles["Heading %d" % lvl]
    st.font.name = "Calibri"
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(_HSP[lvl][0])
    st.paragraph_format.space_after = Pt(_HSP[lvl][1])
    st.paragraph_format.keep_with_next = True

for sec in doc.sections:
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)

# ---- helpers ----
def p(text="", size=10.5, color=INK, bold=False, italic=False, align=None, after=4, before=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(after); par.paragraph_format.space_before = Pt(before)
    if align: par.alignment = align
    r = par.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold; r.italic = italic
    return par

def rich(parts, after=6, align=None):
    """parts = list of (text, bold, color)."""
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if align: par.alignment = align
    for t, b, c in parts:
        r = par.add_run(t); r.bold = b; r.font.color.rgb = c or INK; r.font.size = Pt(10.5)
    return par

def h(text, level=1):
    doc.add_heading(text, level=level)

def bullet(text, bold_lead=None):
    par = doc.add_paragraph(style="List Bullet"); par.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = par.add_run(bold_lead + " "); r.bold = True; r.font.color.rgb = INK; r.font.size = Pt(10.5)
    r2 = par.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = INK
    return par

def step(text):
    par = doc.add_paragraph(style="List Number"); par.paragraph_format.space_after = Pt(3)
    r = par.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = INK
    return par

def shade(par, hexfill):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexfill)
    pPr.append(sh)

def callout(label, text, fill="EEF2FF", color=INDIGO):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(5); par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.left_indent = Inches(0.05)
    r = par.add_run(" " + label + "  "); r.bold = True; r.font.color.rgb = color; r.font.size = Pt(10)
    r2 = par.add_run(text + " "); r2.font.size = Pt(10); r2.font.color.rgb = INK
    shade(par, fill)

_FIG = [0]
def figure(name, caption):
    path = os.path.join(IMG, name)
    if not os.path.exists(path):
        p("[missing image: %s]" % name, color=RED); return
    _FIG[0] += 1
    caption = re.sub(r"^Figure\s+\d+\s+—\s*", "", caption)          # drop any hard-coded number
    caption = "Figure %d — %s" % (_FIG[0], caption)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(1)
    par.paragraph_format.keep_with_next = True
    par.add_run().add_picture(path, width=Inches(6.2))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6); cap.paragraph_format.space_before = Pt(0)
    rc = cap.add_run(caption); rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = MUTE

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(htext); rr.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),"4F46E5")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val)); rr.font.size = Pt(9.5); rr.font.color.rgb = INK
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def page_break():
    pass  # sections now flow continuously; only the cover & TOC force a break (keeps the doc compact)

def _field(par, code, default="1"):
    run = par.add_run()
    def fc(t):
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), t); return e
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = code
    t = OxmlElement("w:t"); t.text = default
    r = run._r
    r.append(fc("begin")); r.append(instr); r.append(fc("separate")); r.append(t); r.append(fc("end"))
    run.font.size = Pt(8.5); run.font.color.rgb = MUTE
    return run

def build_header_footer():
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True  # keep the cover clean
    # header (from page 2)
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("NAZTECH OMS  ·  Exchange-Hosted Order Management System — User Manual")
    hr.font.size = Pt(8.5); hr.font.color.rgb = MUTE
    # footer (from page 2): centered “Page X of Y”
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page "); r1.font.size = Pt(8.5); r1.font.color.rgb = MUTE
    _field(fp, "PAGE", "1")
    r2 = fp.add_run(" of "); r2.font.size = Pt(8.5); r2.font.color.rgb = MUTE
    _field(fp, "NUMPAGES", "1")

def toc():
    par = doc.add_paragraph()
    run = par.add_run()
    def fc(t):
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), t); return e
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    ph = OxmlElement("w:t")
    ph.text = "The contents list builds automatically when you open this document in Word."
    r = run._r
    r.append(fc("begin")); r.append(instr); r.append(fc("separate")); r.append(ph); r.append(fc("end"))

# ============================================================ COVER
build_header_footer()
doc.add_paragraph().paragraph_format.space_after = Pt(60)
p("NAZTECH OMS", size=34, color=INDIGO, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("Exchange-Hosted Order Management System", size=15, color=VIOLET, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("for Dhaka & Chittagong Stock Exchanges (DSE / CSE)", size=12, color=MUTE, align=WD_ALIGN_PARAGRAPH.CENTER, after=40)
p("USER MANUAL", size=20, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
p("A friendly, complete guide for investors, dealers, risk officers and exchange administrators",
  size=11, color=MUTE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
p("License & AMC Model  ·  RFP Ref DSE/CSD/APP/2026", size=10, color=MUTE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("Prepared by Naztech  ·  Version 1.0", size=10, color=MUTE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p("Document is based on the live demonstration build.", size=9, color=MUTE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================ TOC
h("Table of Contents", 1)
toc()
doc.add_page_break()

# ============================================================ 1. INTRODUCTION
h("1. Introduction", 1)

h("1.1  What this product is", 2)
p("Naztech OMS is an Exchange-Hosted Order Management System (OMS) built for Bangladesh’s capital "
  "market — the Dhaka Stock Exchange (DSE) and the Chittagong Stock Exchange (CSE). It is the software "
  "that brokerage houses and their clients use to watch the market, place buy and sell orders, manage "
  "positions and risk, and produce end-of-day reports — all from one modern, browser-based workspace.")
p("An “Order Management System” sits between the people who want to trade (investors and the dealers who "
  "act for them) and the exchange’s Matching Engine (the system that actually pairs buyers with sellers). "
  "The OMS captures each order, checks it against risk rules, sends valid orders to be matched, and tracks "
  "every change to the order through its whole life — from creation to fill, amendment, cancellation or "
  "rejection.")

h("1.2  Purpose & background", 2)
p("This system was developed in response to the DSE tender for the supply, implementation and support of "
  "an Exchange-Hosted OMS under the License & AMC model. The goals were to deliver:")
bullet("a complete multi-asset order lifecycle (equities, bonds, sukuk, mutual funds, derivatives);", "Full trading workflow —")
bullet("pre-trade risk management that protects brokers, clients and the market;", "Built-in risk controls —")
bullet("a pluggable matching gateway that simulates the exchange now and connects to the real Matching Engine later;", "Future-proof design —")
bullet("genuine, on-premise Artificial Intelligence value-adds that work without sending data to the cloud;", "On-prem AI —")
bullet("and a premium, modern user experience that does not look or feel like a traditional grey trading terminal.", "Distinctive UX —")
callout("ABOUT THIS BUILD", "The screens in this manual come from the working demonstration build. Market prices, "
        "orders and trades shown are from a realistic demo data set so that every screen looks full and lifelike. "
        "The same software connects to authentic exchange data feeds and the real Matching Engine in production.")

h("1.3  Who should read this manual", 2)
p("The system serves several kinds of users, each with their own view and permissions. You only need the "
  "sections that match your role, but the manual is written so anyone can read it end to end.")
table(
    ["Role", "Who they are", "What they mainly do"],
    [["Investor (Client)", "A member of the public who owns a trading account", "View own portfolio & P&L, place own orders, ask the AI Advisor"],
     ["Dealer", "Brokerage-house staff who trade for clients", "Place / amend / cancel orders for client accounts, watch the market"],
     ["Trader", "Brokerage dealing-desk operator", "Same trading tools as a dealer, scoped by limits"],
     ["RMS Manager", "Risk & compliance officer at the broker", "Set risk limits, watch the AI risk-alert feed, use the kill-switch"],
     ["Broker Admin", "Brokerage administrator", "Manage the firm’s users and accounts"],
     ["Exchange Admin", "Exchange / CSD control-plane operator", "Onboard brokers, oversee securities, users, orders & audit trail"],
     ["View-only", "Auditors / observers", "Read-only access to screens"]],
    widths=[1.5, 2.4, 2.8])

page_break()

# ============================================================ 2. ARCHITECTURE (light)
h("2. How the system fits together (in plain words)", 1)
p("You do not need any technical knowledge to use the OMS, but a simple mental picture helps. The product "
  "has four cooperating parts:")
bullet("the web application you see and click in your browser (the “front end”);", "1. The Terminal —")
bullet("a core engine that stores data, enforces risk rules, runs the order book and exposes the AI;", "2. The OMS server —")
bullet("a market-data service that brings in live DSE/CSE prices and keeps them fresh;", "3. The Market-Data service —")
bullet("a database that safely keeps every security, order, trade, holding and audit record.", "4. The Database —")
p("Two pieces of intelligence are built in and run entirely on-premise (your data never leaves your "
  "infrastructure): a semantic search that understands what you mean even with typos, and an explainable "
  "risk-scoring model that flags unusual orders. An optional AI Investment Advisor adds conversational, "
  "voice-enabled guidance.")
callout("NOTE", "Because the matching of orders is handled through a “pluggable gateway”, the exact same screens and "
        "workflow you learn here continue to work unchanged when the system is connected to the exchange’s real "
        "Matching Engine.", fill="F0FDF4", color=GREEN)

page_break()

# ============================================================ 3. GETTING STARTED
h("3. Getting started", 1)

h("3.1  Opening the application", 2)
p("The OMS runs in a normal web browser — no software to install on your computer. Use a recent version of "
  "Google Chrome, Microsoft Edge or Mozilla Firefox.")
step("Open your browser.")
step("Go to the address your administrator gives you. On the demo machine this is http://localhost:3060 ; "
     "colleagues on the same network use the host computer’s IP address, for example http://10.33.56.3:3060 .")
step("The sign-in screen appears.")

h("3.2  Signing in", 2)
p("Every user has a single-session, role-based login. Type your username and password and select "
  "“Sign in”. For demonstrations, four one-click “Quick demo logins” are provided on the right of the "
  "sign-in screen so you can instantly experience each role.")
figure("01_login.png", "Figure 1 — The sign-in screen, with one-click Quick Demo Logins for each role.")
p("Demo accounts (all use the password demo123):")
table(
    ["Button / Username", "Role", "Experience"],
    [["investor1", "Investor (Client)", "The customer view — own portfolio, place own orders, AI Advisor"],
     ["dealer1", "Dealer (Broker staff)", "Trade & manage orders on behalf of clients"],
     ["rms", "RMS Manager", "Risk limits, kill-switch and the AI risk-alert feed"],
     ["exchadmin", "Exchange Admin", "The exchange control plane"],
     ["investor1 … investor5", "Investors", "Five different clients, each with a distinct portfolio"]],
    widths=[2.1, 1.8, 2.8])
callout("TIP", "Not sure which login to use? If you are a member of the public who wants to trade your own account, "
        "choose Investor (Client). If you are brokerage staff trading for clients, choose Dealer.")

h("3.3  Signing out", 2)
p("Your name and role are shown at the bottom-left of every screen. Select “Sign out” there to end your "
  "session securely. Because sessions are single-use, signing in again elsewhere will end an earlier session.")

page_break()

# ============================================================ 4. THE INTERFACE
h("4. Getting around the interface", 1)
p("Every screen shares the same simple frame, so once you learn it once you can find your way anywhere.")

h("4.1  The left navigation menu", 2)
p("The dark sidebar on the left is your main menu. Select any item to open that module. The item you are "
  "currently on is highlighted.")
table(
    ["Menu item", "What it opens"],
    [["My Dashboard", "Your personal, customisable board of charts and widgets"],
     ["Trader Terminal", "The full trading screen — watch, chart, order ticket, blotter"],
     ["Trading Desk", "A dockable, re-arrangeable version of the trading screen"],
     ["Market Screener", "A searchable, filterable table of every instrument"],
     ["Portfolio", "Your holdings, profit & loss and allocation analytics"],
     ["Reports", "Trade book, order book, position statement and EOD export"],
     ["Risk (RMS)", "Risk limits, broker kill-switch and AI alert feed (RMS/Admin)"],
     ["Exchange Admin", "The exchange control plane (Exchange Admin only)"]],
    widths=[1.8, 4.9])
p("Some items only appear for the roles that may use them — an Investor, for example, focuses on Dashboard, "
  "Terminal, Screener and Portfolio.")

h("4.2  The top bar", 2)
p("Across the top of every screen you will find:")
bullet("the page title on the left;", "Title —")
bullet("page-specific controls such as the AI search box or the DSE/CSE exchange switch;", "Context tools —")
bullet("the Theme switcher to change the colour scheme (see 4.4);", "Theme —")
bullet("a Live / Offline indicator — a green pulsing dot means real-time updates are streaming in;", "Live status —")
bullet("a running market clock in Bangladesh Standard Time.", "Clock —")

h("4.3  The AI Advisor button", 2)
p("A floating “AI Advisor” button sits at the bottom-right of every screen. Select it any time to open the "
  "conversational assistant (covered fully in section 5.9).")

h("4.4  Choosing a theme", 2)
p("Select the theme switch in the top bar to pick from five looks: Midnight (the dark default), Daylight "
  "(light), Terminal (retro green-amber), Ocean (cool blue) and Sunset (warm). Your choice is remembered on "
  "your device. This manual’s screenshots use the Midnight theme.")

h("4.5  Moving, resizing & docking panels", 2)
rich([("The interface is yours to arrange. There are two ways to customise it:", False, INK)])
bullet("drag a widget by its title bar to move it; drag its bottom-right corner to resize; use the small "
       "icons on each widget to minimise, maximise or hide it. Your layout is saved automatically.",
       "On My Dashboard —")
bullet("drag any panel’s tab to re-dock it (left, right, top, bottom or as a new tab); drag the borders to "
       "resize; use the ⛶ icon to maximise a panel and ✕ to close it. “+ Add panel” brings panels back and "
       "“Reset layout” restores the default.", "On the Trading Desk —")

h("4.6  Quick actions worth knowing", 2)
bullet("type a symbol, a company name, or even a plain-language phrase (with typos) — the AI finds the right "
       "instrument and lets you jump straight to trading it.", "AI search —")
bullet("in the order book (market depth), click any price to copy it straight into the order ticket.", "Click-to-price —")
bullet("press Ctrl + V inside the AI Advisor to paste a screenshot of a chart for the AI to read.", "Paste a chart —")
bullet("use the DSE / CSE switch in the top bar to flip the whole screen between the two exchanges.", "Switch exchange —")

page_break()

# ============================================================ 5. MODULES
h("5. The modules, one by one", 1)
p("Each module below opens with what it is for, then a short how-to. Screens vary slightly by role and "
  "permissions.")

# 5.1 Dashboard
h("5.1  My Dashboard", 2)
p("Purpose: your personal command centre. Build a board of exactly the charts and panels you care about — "
  "market breadth, sector performance, top gainers and losers, candlestick charts, a heatmap, your P&L over "
  "time, holdings and more — then arrange them however you like.")
figure("02_dashboard.png", "Figure 2 — My Dashboard with the “Market Overview” preset: index board, market-breadth "
       "donut, sector performance, gainers/losers, most-active and a candlestick chart.")
p("How to use it:")
step("Select “+ Add widget” (top right) and pick widgets from the menu, grouped by Market, Portfolio, "
     "Orders & Risk and News.")
step("Drag each widget by its title to position it; drag the corner to resize; use the icons to minimise, "
     "maximise or hide it.")
step("Use “Load preset…” to instantly apply a ready-made board (Market Overview, Trading Desk, Portfolio & "
     "P&L, Risk & Compliance), or select “Save” to store your own named layout.")
step("Select “Reset” to return to the default board. Your layout is remembered automatically per user.")

# 5.2 Terminal
h("5.2  Trader Terminal", 2)
p("Purpose: the heart of day-to-day trading. Everything you need to trade one instrument is on a single "
  "screen — a live market watch, a price-and-volume chart, the order book (market depth), the order ticket, "
  "your order blotter, your portfolio and the latest news.")
figure("03_terminal.png", "Figure 3 — The Trader Terminal: market watch (left), chart & order book (centre), "
       "order ticket with AI pre-trade risk (right), and the order blotter along the bottom.")
p("Placing an order:")
step("Find your instrument — click it in the Market Watch list, or use the AI search box at the top.")
step("In the Order Ticket (right), choose BUY or SELL.")
step("Set Type (Limit, Market, Stop, Stop-Limit), Window (Normal, Spot, Block, Odd-lot…), Quantity, Price "
     "and Validity (Day, GTC, GTD…). Tip: click a price in the order book to fill it in automatically.")
step("Watch the AI Pre-Trade Risk meter — it scores the order and explains any concern before you send.")
step("Select the green “BUY/SELL …” button. Confirm the account at the foot of the ticket first.")
p("Tracking & changing orders:")
bullet("every order you place appears in the Order Blotter at the bottom with its live status (Open, "
       "Partial, Filled, Cancelled, Rejected).", "Blotter —")
bullet("use the ✎ (amend) action on an open order to change its price or quantity; use the cancel action to "
       "withdraw it.", "Amend / cancel —")
bullet("build and switch named Watchlists above the market list, and filter to Gainers or Losers in one click.",
       "Watchlists —")
figure("14_watchlist.png", "Figure 4 — Watchlists close-up: use the selector to switch between All Securities, "
       "Gainers, Losers or your own saved lists; “+” creates a new watchlist and ★ adds the selected instrument.")
figure("13_amend.png", "Figure 5 — Amending an order: the ✎ action on an open order opens the Amend Order "
       "dialog, where you can change the price or quantity and re-submit.")
p("Understanding an order’s status: every order moves through a clear life-cycle, always visible in the "
  "blotter and the audit trail.")
table(
    ["Status", "What it means"],
    [["New / Pending Risk", "Just submitted and being risk-checked"],
     ["Open", "Accepted and resting in the order book, waiting to match"],
     ["Partial", "Partly filled; the remainder is still working"],
     ["Filled", "Completely executed"],
     ["Cancelled", "Withdrawn by the user before fully filling"],
     ["Rejected", "Blocked by a risk rule (the reason is shown on the order)"],
     ["Expired", "Validity period ended before it filled"]],
    widths=[1.7, 5.1])

h("Reading the Market Depth (order book)", 3)
p("The Market Depth panel — shown below the chart on the Trader Terminal, and as the dockable "
  "“Order Book (Depth)” panel on the Trading Desk — displays the live order book: the buy (bid) and sell "
  "(ask) orders waiting at each price level. It is the clearest picture of current supply, demand and likely "
  "short-term price movement, and is where most traders decide their price.")
figure("16_depth_crop.png", "Market Depth (order book): bid levels in green on the left and ask levels in red "
       "on the right, each with its quantity and order count, around the Last Traded Price (LTP).")
p("How to read it:")
bullet("the green BID side (left) lists buyers and the prices they are willing to pay — best (highest) bid at the top.", "Bids —")
bullet("the red ASK side (right) lists sellers and the prices they want — best (lowest) ask at the top.", "Asks —")
bullet("QTY is the total quantity available at that price; ORD is how many separate orders make it up; the "
       "shaded bar shows each level’s size at a glance.", "Quantity & orders —")
bullet("LTP at the top-right is the Last Traded Price; the gap between the best bid and best ask is the spread.", "LTP & spread —")
bullet("click any price in the ladder to copy it straight into the Order Ticket — the quickest way to price an order.", "Click-to-price —")

# 5.3 Trading Desk
h("5.3  Trading Desk (dockable workspace)", 2)
p("Purpose: the same trading tools as the Terminal, but as a professional, fully dockable workspace — like "
  "the multi-panel desks used on institutional trading floors. Power users can arrange panels exactly to "
  "their preference and the layout is remembered.")
figure("04_workspace.png", "Figure 4 — The Trading Desk: every panel (Market Watch, Chart, Order Book, Order "
       "Ticket, Portfolio, News, Blotter) can be dragged, split, resized, maximised or closed.")
p("How to arrange it:")
step("Drag any panel’s tab and drop it to the left, right, top or bottom of another panel to split the "
     "space — or drop it onto a tab strip to stack it as a new tab.")
step("Drag the borders between panels to resize them.")
step("Use the ⛶ icon to maximise a panel to full screen, and ✕ to close one.")
step("“+ Add panel” brings any closed panel back; “Reset layout” restores the default arrangement.")
callout("TIP", "The Trading Desk and the classic Trader Terminal share the same live data and account selection, "
        "so you can use whichever feels more comfortable.")

# 5.4 Screener
h("5.4  Market Screener", 2)
p("Purpose: scan the whole market quickly. A searchable, sortable table of every instrument with last price, "
  "change, volume, turnover, day high/low and more — with filters to narrow down to exactly what you want.")
figure("05_screener.png", "Figure 5 — The Market Screener: filter by asset class, sector, movers and price band, "
       "then sort any column or jump straight to trading.")
p("How to use it:")
step("Type in the search box to find by symbol or company name (the AI tolerates typos).")
step("Use the drop-downs to filter by asset class, sector or movers (gainers/losers), and set a min/max "
     "price band.")
step("Click any column header to sort by it.")
step("Select ★ to add an instrument to a watchlist, or “Trade ↗” to open it in the Terminal ready to order.")

# 5.5 Portfolio
h("5.5  Portfolio", 2)
p("Purpose: understand what you own and how it is performing. See cash, market value, realised and "
  "unrealised profit & loss, your equity curve over time, allocation by sector and asset class, and a full "
  "holdings table.")
figure("06_portfolio.png", "Figure 6 — The Portfolio page: headline P&L tiles, an equity curve over time, "
       "allocation donuts and a per-holding profit & loss table.")
p("How to read it:")
bullet("the tiles across the top summarise cash, invested value, market value, day P&L and unrealised P&L.",
       "Headline numbers —")
bullet("the equity curve shows how your account value (or P&L) has moved over time — toggle between the two.",
       "P&L over time —")
bullet("the donuts show how your money is spread across sectors and asset classes.", "Allocation —")
bullet("the holdings table lists each position with quantity, average cost, last price, value, day and "
       "unrealised P&L, and portfolio weight.", "Holdings —")

# 5.6 Reports
h("5.6  Reports & End-of-Day export", 2)
p("Purpose: produce the records you need for compliance, settlement and your own records — a Trade Book, an "
  "Order Book, a Position Statement and an EOD Summary — and export them.")
figure("07_reports.png", "Figure 7 — The Reports module showing the Trade Book, with one-click CSV and "
       "Print / PDF export.")
p("How to use it:")
step("Choose a report tab: Trade Book, Order Book, Position Statement or EOD Summary.")
step("For the Position Statement, pick the account from the selector.")
step("Select “⤓ CSV” to download a spreadsheet, or “⎙ Print / PDF” to print or save a clean PDF of the "
     "report.")

# 5.7 RMS
h("5.7  Risk Management (RMS)", 2)
p("Purpose: keep trading inside safe limits and react instantly to trouble. The RMS module (for RMS "
  "Managers and Exchange Admins) sets the pre-trade risk limits applied to every order, provides a per-broker "
  "kill-switch, and shows a live, AI-powered risk-alert feed.")
figure("08_rms.png", "Figure 8 — Risk Management: the broker kill-switch (top), editable pre-trade limits "
       "(centre), the controls enforced at order entry, and the live AI risk-alert feed (right).")
p("What you can do:")
bullet("instantly halt or resume all new orders for a broker with the HALT / RESUME button — used in an "
       "emergency.", "Kill-switch —")
bullet("set the maximum order value, maximum quantity, mark-to-market loss limit, wash-sale block and on/off "
       "switch per scope (Broker, Trader, Client), then select Save. Changes apply live at the next order.",
       "Editable limits —")
bullet("every order is checked against firm/trader/client limits, buying-power, holdings (short-sell), "
       "wash-sale, lot-size, stop-loss, AI fat-finger detection and the kill-switch.", "Controls —")
bullet("watch flagged orders stream in with an explainable risk score and reason (for example a rejected "
       "wash-sale).", "AI risk-alert feed —")

# 5.8 Admin
h("5.8  Exchange Admin (control plane)", 2)
p("Purpose: the exchange / CSD operator’s overview and controls — onboard brokers (TREC holders), see the "
  "live counts of brokers, users, securities, orders, trades and connected clients, review the user hierarchy "
  "and read the system audit trail.")
figure("09_admin.png", "Figure 9 — The Exchange Control Plane: system totals, the on-prem AI engine status, "
       "TREC-holder (broker) list, user hierarchy, audit trail, and the broker-onboarding form.")
p("Highlights:")
bullet("at-a-glance totals for brokers, users, securities, orders, trades and live clients.", "Dashboard tiles —")
bullet("complete the TREC code, firm name, firm limit and exchange, then select “Onboard broker”.",
       "Onboard a broker —")
bullet("see every user by role across the system.", "User hierarchy —")
bullet("an append-only record of logins and key actions for compliance.", "Audit trail —")

# 5.9 AI Advisor
h("5.9  AI Investment Advisor", 2)
p("Purpose: a friendly, conversational assistant grounded in live OMS data. Ask it about any DSE/CSE stock, "
  "your portfolio and P&L, sectors, share categories or order types — by typing or by speaking, in English or "
  "Bangla — and it can reply in chat and, if you wish, read the answer aloud.")
figure("12_ai_advisor.png", "Figure 10 — The AI Investment Advisor panel, with suggestion chips, a DSE Status "
       "button, English/Bangla toggle, and voice & screenshot input.")
p("How to use it:")
step("Select the “AI Advisor” button (bottom-right of any screen).")
step("Type your question, or select the microphone and speak (English or Bangla). You can also press "
     "Ctrl + V to paste a screenshot of a chart for the AI to read.")
step("Use the language toggle (EN / বাং) to set the reply language, and the suggestion chips for quick "
     "questions.")
step("Select “DSE Status” for an instant market snapshot, or the 🔊 Listen button on any answer to hear it "
     "spoken.")
callout("PLEASE NOTE", "The AI Advisor is informational only and is not licensed financial advice. Always make your "
        "own decisions or consult a licensed adviser.", fill="FEF2F2", color=RED)

page_break()

# ============================================================ 6. ROLE WALKTHROUGHS
h("6. Quick walkthroughs by role", 1)

h("6.1  I am an Investor (the public)", 2)
p("Sign in as investor1 / demo123. You land on your Dashboard, scoped to your own account.")
step("Open Portfolio to see what you own and your profit & loss.")
step("Open the Trader Terminal, find a stock, and place a buy or sell order for your own account.")
step("Open the AI Advisor and ask, for example, “Am I in profit or loss?” or “Is GP a good position today?”.")
figure("11_investor_portfolio.png", "Figure 11 — An investor’s own Portfolio view (Ayesha Rahman), with equity "
       "curve, allocation and holdings.")

h("6.2  I am a Dealer (broker staff)", 2)
step("Sign in as dealer1 / demo123.")
step("Choose the client account at the foot of the Order Ticket, then place, amend or cancel orders for "
     "that client.")
step("Use the Trading Desk if you prefer a dockable, multi-panel desk.")
step("Use Reports at end of day to export the trade book and position statements.")

h("6.3  I am an RMS Manager", 2)
step("Sign in as rms / demo123 — you land on the Risk (RMS) screen.")
step("Review and adjust the pre-trade limits; select Save to apply them live.")
step("Watch the AI risk-alert feed; if something is wrong, use the broker kill-switch to HALT new orders.")

h("6.4  I am an Exchange Admin", 2)
step("Sign in as exchadmin / demo123 — you land on the Exchange Control Plane.")
step("Onboard brokers, review system totals, the user hierarchy and the audit trail.")
figure("10_investor_dashboard.png", "Figure 12 — A role-tailored dashboard (here an investor’s “Portfolio & "
       "P&L” preset) — every user can shape their own board.")

page_break()

# ============================================================ 7. COMMON TASKS
h("7. Common tasks at a glance", 1)
table(
    ["I want to…", "Do this"],
    [["Place an order", "Terminal → pick instrument → Order Ticket → set side/type/qty/price → BUY/SELL"],
     ["Change an order", "Terminal → Order Blotter → ✎ on the open order → edit price/qty → save"],
     ["Cancel an order", "Terminal → Order Blotter → cancel action on the open order"],
     ["Make a watchlist", "Terminal → watchlist bar → “+” → name it → add instruments with ★"],
     ["Find a stock fast", "Use the AI search box (symbol, company name, or plain words — typos OK)"],
     ["See my P&L", "Portfolio → headline tiles, equity curve and holdings table"],
     ["Export a report", "Reports → choose tab → ⤓ CSV or ⎙ Print / PDF"],
     ["Build my dashboard", "My Dashboard → + Add widget → drag/resize → Save preset"],
     ["Arrange my desk", "Trading Desk → drag tabs/borders → ⛶ maximise → Reset layout"],
     ["Change limits", "Risk (RMS) → edit a limit row → Save (applies live)"],
     ["Emergency stop", "Risk (RMS) → broker kill-switch → HALT"],
     ["Switch exchange", "Top bar → DSE / CSE toggle"],
     ["Change the look", "Top bar → Theme switch → pick a theme"],
     ["Ask the AI", "AI Advisor button → type or speak (EN/বাং)"]],
    widths=[2.2, 4.6])

# ============================================================ 8. GLOSSARY
h("8. Glossary", 1)
table(
    ["Term", "Meaning"],
    [["OMS", "Order Management System — captures, risk-checks, routes and tracks orders"],
     ["Matching Engine", "The exchange system that pairs buy and sell orders into trades"],
     ["TREC", "Trading Right Entitlement Certificate — a broker’s right to trade on the exchange"],
     ["RMS", "Risk Management System — the pre-trade and monitoring risk controls"],
     ["LTP", "Last Traded Price — the most recent price an instrument traded at"],
     ["Market depth / Order book", "The list of resting buy and sell orders at each price level"],
     ["Limit order", "An order with a maximum buy / minimum sell price"],
     ["Market order", "An order to trade immediately at the best available price"],
     ["Validity (Day/GTC/GTD)", "How long an order stays active (today / till cancelled / till a date)"],
     ["Wash sale", "Buying and selling the same security to create misleading activity — blocked by RMS"],
     ["Lot size", "The minimum tradable quantity for an instrument"],
     ["P&L", "Profit & Loss — realised (booked) and unrealised (on open positions)"],
     ["Share categories A/B/N/Z", "DSE categories by dividend & compliance record (Z = highest caution)"],
     ["Kill-switch", "An RMS control that instantly halts all new orders for a broker"]],
    widths=[2.0, 4.8])

# ============================================================ 9. TROUBLESHOOTING
h("9. Troubleshooting & FAQ", 1)
def qa(q, a):
    rich([(q + "  ", True, INDIGO), (a, False, INK)])
qa("The top bar shows “Offline”.",
   "Real-time streaming is briefly disconnected. The screen still works; it usually reconnects on its own. "
   "If it persists, refresh the page.")
qa("A colleague on the network cannot sign in.",
   "Make sure they use the host computer’s IP address (e.g. http://10.33.56.3:3060), not “localhost”, and "
   "that the host’s firewall allows the OMS ports.")
qa("The dashboard is empty.",
   "Select “+ Add widget” to add panels, or “Load preset…” to apply a ready-made board.")
qa("My order was rejected.",
   "Open Risk (RMS) or read the reason shown on the order in the blotter — it will name the limit or check "
   "that blocked it (for example a wash-sale or an order-value limit).")
qa("The AI Advisor answered in the wrong language.",
   "Use the EN / বাং toggle at the top of the Advisor panel to set the reply language before asking.")
qa("Can I change which account I trade?",
   "Dealers and admins can pick the account at the foot of the Order Ticket. Investors are tied to their own "
   "account.")

doc.add_paragraph()
p("— End of manual —", color=MUTE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# tell Word to refresh all fields (the TOC) when the document is opened
_upd = OxmlElement("w:updateFields"); _upd.set(qn("w:val"), "true")
doc.settings.element.append(_upd)

doc.save(OUT)
print("Saved:", OUT, "| figures:", _FIG[0])
